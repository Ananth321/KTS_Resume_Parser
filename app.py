# import base64
# import zipfile
# from flask import Flask, render_template, request, send_file, redirect, url_for, flash, session
# from werkzeug.utils import secure_filename
# from docx import Document
# import os
# import io
# from datetime import datetime
# import re
# import docx2txt
# import glob
# import uuid
# from docx.shared import Inches
# from PIL import Image
# import io
#
# app = Flask(__name__)
# app.secret_key = 'your-secret-key-here'
# app.config['MAX_CONTENT_LENGTH'] = 2 * 1024 * 1024  # 2MB limit
# app.config['UPLOAD_FOLDER'] = 'uploads'
# app.config['ALLOWED_EXTENSIONS'] = {'docx'}
# app.config['PHOTO_EXTRACTION'] = True
#
# os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
#
# def allowed_file(filename):
#     return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']
#
#
# def extract_contact_info(text):
#     contact = {}
#
#     # Improved email regex
#     email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b', text)
#     if email_match:
#         contact['email'] = email_match.group(0)
#
#     # Improved phone regex for international numbers
#     phone_match = re.search(r'\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', text)
#     if phone_match:
#         contact['phone'] = phone_match.group(0)
#
#     # Improved LinkedIn URL regex
#     linkedin_match = re.search(
#         r'(?:(?:http|https)://)?(?:www\.)?linkedin\.com/(?:in|pub|profile)/[a-zA-Z0-9-]+/?\b',
#         text
#     )
#     if linkedin_match:
#         linkedin_url = linkedin_match.group(0)
#         if not linkedin_url.startswith('http'):
#             linkedin_url = 'https://' + linkedin_url
#         contact['linkedin'] = linkedin_url.rstrip('/')
#
#     return contact
#
# def extract_name(text):
#     excluded_headers = {
#         'resume', 'cv', 'curriculum vitae', 'personal profile',
#         'contact information', 'professional profile','career summary'
#     }
#
#     name_patterns = [
#         r'^[A-Z][a-z]+ [A-Z][a-z]+$',
#         r'^[A-Z][a-z]+ [A-Z]\. [A-Z][a-z]+$',
#         r'^[A-Z][a-z]+ [A-Z][a-z]+ [A-Z][a-z]+$',
#         r'^[A-Z][a-z]+, [A-Z][a-z]+$',
#     ]
#
#     lines = [line.strip() for line in text.split('\n') if line.strip()]
#     for line in lines[:5]:
#         if line.lower() in excluded_headers or '@' in line or 'phone' in line.lower():
#             continue
#
#         for pattern in name_patterns:
#             if re.fullmatch(pattern, line):
#                 return line
#
#         words = line.split()
#         if 2 <= len(words) <= 4 and all(word.istitle() for word in words):
#             if not any(word.lower() in {'linkedin', 'github', 'portfolio'} for word in words):
#                 return line
#
#     for line in lines[:5]:
#         if len(line.split()) >= 2:
#             return line
#
#     return None
# def extract_photo_from_docx(filepath):
#     MIN_WH_INCH   = 1
#     MIN_FILE_SIZE = 20 * 1024
#
#     min_wh_emus = Inches(MIN_WH_INCH)
#     doc         = Document(filepath)
#
#     def save_and_return(blob, orig_name):
#         if len(blob) < MIN_FILE_SIZE:
#             return None
#         photo_id  = str(uuid.uuid4())
#         ext       = os.path.splitext(orig_name)[1].lower()
#         fname     = f"photo_{photo_id}{ext}"
#         out_path  = os.path.join(app.config['UPLOAD_FOLDER'], fname)
#         with open(out_path, "wb") as f:
#             f.write(blob)
#         return photo_id
#     for shp in doc.inline_shapes:
#         if shp.width < min_wh_emus or shp.height < min_wh_emus:
#             continue
#         blip      = shp._inline.graphic.graphicData.pic.blipFill.blip
#         rel_id    = blip.embed
#         part      = doc.part.related_parts[rel_id]
#         blob      = part.blob
#         pid = save_and_return(blob, part.filename)
#         if pid:
#             return pid
#     for rel in doc.part.rels.values():
#         if rel.reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image":
#             part = rel._target
#             blob = part.blob
#             try:
#                 elm  = rel._target_part._blob
#             except Exception:
#                 pass
#             pid = save_and_return(blob, part.filename)
#             if pid:
#                 return pid
#     try:
#         with zipfile.ZipFile(filepath) as zf:
#             for name in zf.namelist():
#                 if not name.startswith("word/media/"):
#                     continue
#                 if not name.lower().endswith((".jpg", ".jpeg", ".png")):
#                     continue
#                 blob = zf.read(name)
#                 pid  = save_and_return(blob, name)
#                 if pid:
#                     return pid
#     except Exception:
#         pass
#     return  None
#
#
#
# def parse_resume(filepath):
#     text = docx2txt.process(filepath)
#     resume_data = {
#         'contact': {},
#         'summary': '',
#         'education': [],
#         'experience': [],
#         'skills': [],
#         'projects': [],
#         'achievements': [],
#         'certifications': [],
#         'photo': None
#     }
#
#     if app.config['PHOTO_EXTRACTION']:
#         resume_data['photo'] = extract_photo_from_docx(filepath)
#
#     name = extract_name(text)
#     if name:
#         resume_data['contact']['name'] = name
#
#     resume_data['contact'].update(extract_contact_info(text))
#
#     sections = re.split(r'\n\s*\n', text)
#     current_section = None
#
#     section_patterns = [
#         ('summary', r'^(profile\s*summary|summary|about|objective|career\s*objectives|professional\s*summary|career\s*summary|profile)'),
#         ('experience', r'^(work\s*experience|experience|professional\s*experience|internship)'),
#         ('education', r'^(education|academic\s*background|qualifications|academics)'),
#         ('skills', r'^(skills|technical\s*skills|competencies|key\s*skills|technical\s*expertise)'),
#         ('projects', r'^(projects|key\s*projects|p r o j e c t)'),
#         ('achievements', r'^(achievements|awards|honors|publications)'),
#         ('certifications', r'^(certifications|licenses|courses|certified\s*courses|C E R T I F I C A T I O N)')
#     ]
#
#     for content in sections:
#         content = content.strip()
#         if not content:
#             continue
#
#         section_found = False
#         for section_name, pattern in section_patterns:
#             if re.match(pattern, content, re.IGNORECASE):
#                 current_section = section_name
#                 resume_data[current_section] = [] if isinstance(resume_data[current_section], list) else ''
#                 section_found = True
#                 break
#
#         if not section_found and current_section:
#             if current_section == 'summary':
#                 resume_data[current_section] += ' ' + content
#             elif current_section in ['skills', 'certifications']:
#                 items = re.split(r'[,•;\n]', content)
#                 resume_data[current_section].extend([item.strip() for item in items if item.strip()])
#             else:
#                 lines = [e.strip() for e in content.split('\n') if e.strip()]
#                 resume_data[current_section].extend(lines)
#
#     if resume_data['summary']:
#         resume_data['summary'] = re.sub(
#             r'^(profile\s*summary|summary|about|objective|profile|career\s*objectives)[:\s-]*',
#             '',
#             resume_data['summary'],
#             flags=re.IGNORECASE
#         ).strip()
#
#     for section in list(resume_data.keys()):
#         if isinstance(resume_data[section], (list, dict)) and not resume_data[section]:
#             del resume_data[section]
#         elif isinstance(resume_data[section], str) and not resume_data[section]:
#             del resume_data[section]
#
#     return resume_data
#
# def create_output_doc(sections, original_filename):
#     doc = Document()
#
#     doc.add_heading('Parsed Resume Information', level=1)
#     doc.add_paragraph(f"Original file: {original_filename}")
#     doc.add_paragraph(f"Parsed on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
#
#     if 'photo_id' in session and session['photo_id']:
#         try:
#             for ext in ['.jpg', '.jpeg', '.png']:
#                 photo_path = os.path.join(app.config['UPLOAD_FOLDER'], f"photo_{session['photo_id']}{ext}")
#                 if os.path.exists(photo_path):
#                     doc.add_heading('Profile Photo', level=2)
#                     paragraph = doc.add_paragraph()
#                     run = paragraph.add_run()
#                     run.add_picture(photo_path, width=Inches(1.5))
#                     break
#         except Exception as e:
#             app.logger.error(f"Error adding photo to document: {str(e)}")
#
#     if 'contact' in sections:
#         doc.add_heading('Contact Information', level=2)
#         contact_table = doc.add_table(rows=0, cols=2)
#         for field in ['name', 'email', 'phone', 'address', 'linkedin']:
#             if sections['contact'].get(field):
#                 row = contact_table.add_row()
#                 row.cells[0].text = field.capitalize() + ':'
#                 row.cells[1].text = sections['contact'][field]
#
#     if sections.get('summary'):
#         doc.add_heading('Profile Summary', level=2)
#         doc.add_paragraph(sections['summary'])
#
#     for sec in ['skills', 'experience', 'education', 'certifications', 'projects', 'achievements']:
#         if sections.get(sec):
#             doc.add_heading(sec.capitalize(), level=2)
#             for item in sections[sec]:
#                 doc.add_paragraph(item, style='List Bullet')
#
#     output_stream = io.BytesIO()
#     doc.save(output_stream)
#     output_stream.seek(0)
#     return output_stream
#
# @app.route('/', methods=['GET', 'POST'])
# def upload_file():
#     if request.method == 'POST':
#         if 'file' not in request.files or not request.files['file']:
#             flash('No file selected')
#             return redirect(request.url)
#
#         file = request.files['file']
#
#         if file.filename == '':
#             flash('No file selected')
#             return redirect(request.url)
#
#         if file and allowed_file(file.filename):
#             try:
#                 filename = secure_filename(file.filename)
#                 filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#                 file.save(filepath)
#
#                 parsed_data = parse_resume(filepath)
#
#                 photo_id = parsed_data.pop('photo', None)
#                 session['parsed_data'] = parsed_data
#                 session['original_filename'] = filename
#                 session['photo_id'] = photo_id
#
#                 os.remove(filepath)
#                 return redirect(url_for('preview'))
#
#             except Exception as e:
#                 flash(f'Error processing file: {str(e)}')
#                 return redirect(request.url)
#         else:
#             flash('Only .docx files are allowed (max 2MB)')
#             return redirect(request.url)
#
#     return render_template('upload.html')
#
# @app.route('/photo/<photo_id>')
# def serve_photo(photo_id):
#     try:
#         for ext in ['.jpg', '.jpeg', '.png']:
#             photo_path = os.path.join(app.config['UPLOAD_FOLDER'], f"photo_{photo_id}{ext}")
#             if os.path.exists(photo_path):
#                 return send_file(photo_path, mimetype=f'image/{ext[1:]}')
#         return "Photo not found", 404
#     except Exception as e:
#         return "Error loading photo", 500
#
# @app.route('/preview')
# def preview():
#     if 'parsed_data' not in session:
#         flash('No resume to preview. Please upload a file first.')
#         return redirect(url_for('upload_file'))
#
#     return render_template('preview.html',
#                            parsed_text=session['parsed_data'],
#                            original_filename=session['original_filename'])
#
# @app.route('/download')
# def download_file():
#     if 'parsed_data' not in session:
#         flash('No resume to download. Please upload a file first.')
#         return redirect(url_for('upload_file'))
#
#     try:
#         output = create_output_doc(
#             session['parsed_data'],
#             session['original_filename']
#         )
#
#         return send_file(
#             output,
#             as_attachment=True,
#             download_name=f"parsed_{session['original_filename']}",
#             mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
#         )
#     except Exception as e:
#         flash(f'Error generating document: {str(e)}')
#         return redirect(url_for('preview'))
#
# if __name__ == '__main__':
#     app.run(debug=True)
#
#
# <====================================for zip file we use this code==========================================================>

import os
import uuid
import zipfile
import shutil
from flask import Flask, render_template, request, send_file, redirect, url_for, flash, session
from werkzeug.utils import secure_filename
from docx import Document
import io
from datetime import datetime
import re
import docx2txt
from docx.shared import Inches

app = Flask(__name__)
app.secret_key = 'your-secret-key-here'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB limit
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'docx', 'zip'}
app.config['PHOTO_EXTRACTION'] = True

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


def cleanup_uploads():
    # Clean up old files in uploads folder
    for filename in os.listdir(app.config['UPLOAD_FOLDER']):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print(f'Failed to delete {file_path}. Reason: {e}')


def extract_contact_info(text):
    contact = {}
    email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b', text)
    if email_match:
        contact['email'] = email_match.group(0)

    phone_match = re.search(r'\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', text)
    if phone_match:
        contact['phone'] = phone_match.group(0)

    linkedin_match = re.search(
        r'(?:(?:http|https)://)?(?:www\.)?linkedin\.com/(?:in|pub|profile)/[a-zA-Z0-9-]+/?\b',
        text
    )
    if linkedin_match:
        linkedin_url = linkedin_match.group(0)
        if not linkedin_url.startswith('http'):
            linkedin_url = 'https://' + linkedin_url
        contact['linkedin'] = linkedin_url.rstrip('/')

    return contact


def extract_name(text):
    excluded_headers = {
        'resume', 'cv', 'curriculum vitae', 'personal profile',
        'contact information', 'professional profile', 'career summary'
    }

    name_patterns = [
        r'^[A-Z][a-z]+ [A-Z][a-z]+$',
        r'^[A-Z][a-z]+ [A-Z]\. [A-Z][a-z]+$',
        r'^[A-Z][a-z]+ [A-Z][a-z]+ [A-Z][a-z]+$',
        r'^[A-Z][a-z]+, [A-Z][a-z]+$',
    ]

    lines = [line.strip() for line in text.split('\n') if line.strip()]
    for line in lines[:5]:
        if line.lower() in excluded_headers or '@' in line or 'phone' in line.lower():
            continue

        for pattern in name_patterns:
            if re.fullmatch(pattern, line):
                return line

        words = line.split()
        if 2 <= len(words) <= 4 and all(word.istitle() for word in words):
            if not any(word.lower() in {'linkedin', 'github', 'portfolio'} for word in words):
                return line

    for line in lines[:5]:
        if len(line.split()) >= 2:
            return line

    return None


def extract_photo_from_docx(filepath):
    MIN_WH_INCH = 1
    MIN_FILE_SIZE = 20 * 1024

    min_wh_emus = Inches(MIN_WH_INCH)
    doc = Document(filepath)

    def save_and_return(blob, orig_name):
        if len(blob) < MIN_FILE_SIZE:
            return None
        photo_id = str(uuid.uuid4())
        ext = os.path.splitext(orig_name)[1].lower()
        fname = f"photo_{photo_id}{ext}"
        out_path = os.path.join(app.config['UPLOAD_FOLDER'], fname)
        with open(out_path, "wb") as f:
            f.write(blob)
        return photo_id

    for shp in doc.inline_shapes:
        if shp.width < min_wh_emus or shp.height < min_wh_emus:
            continue
        blip = shp._inline.graphic.graphicData.pic.blipFill.blip
        rel_id = blip.embed
        part = doc.part.related_parts[rel_id]
        blob = part.blob
        pid = save_and_return(blob, part.filename)
        if pid:
            return pid

    for rel in doc.part.rels.values():
        if rel.reltype == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image":
            part = rel._target
            blob = part.blob
            try:
                elm = rel._target_part._blob
            except Exception:
                pass
            pid = save_and_return(blob, part.filename)
            if pid:
                return pid

    try:
        with zipfile.ZipFile(filepath) as zf:
            for name in zf.namelist():
                if not name.startswith("word/media/"):
                    continue
                if not name.lower().endswith((".jpg", ".jpeg", ".png")):
                    continue
                blob = zf.read(name)
                pid = save_and_return(blob, name)
                if pid:
                    return pid
    except Exception:
        pass
    return None


def parse_resume(filepath):
    text = docx2txt.process(filepath)
    resume_data = {
        'contact': {},
        'summary': '',
        'education': [],
        'experience': [],
        'skills': [],
        'projects': [],
        'achievements': [],
        'certifications': [],
        'photo': None
    }

    if app.config['PHOTO_EXTRACTION']:
        resume_data['photo'] = extract_photo_from_docx(filepath)

    name = extract_name(text)
    if name:
        resume_data['contact']['name'] = name

    resume_data['contact'].update(extract_contact_info(text))

    sections = re.split(r'\n\s*\n', text)
    current_section = None

    section_patterns = [
        ('summary',
         r'^(profile\s*summary|summary|about|objective|career\s*objectives|professional\s*summary|career\s*summary|profile)'),
        ('experience', r'^(work\s*experience|experience|professional\s*experience|internship)'),
        ('education', r'^(education|academic\s*background|qualifications|academics)'),
        ('skills', r'^(skills|technical\s*skills|competencies|key\s*skills|technical\s*expertise)'),
        ('projects', r'^(projects|key\s*projects|p r o j e c t)'),
        ('achievements', r'^(achievements|awards|honors|publications)'),
        ('certifications', r'^(certifications|licenses|courses|certified\s*courses|C E R T I F I C A T I O N)')
    ]

    for content in sections:
        content = content.strip()
        if not content:
            continue

        section_found = False
        for section_name, pattern in section_patterns:
            if re.match(pattern, content, re.IGNORECASE):
                current_section = section_name
                resume_data[current_section] = [] if isinstance(resume_data[current_section], list) else ''
                section_found = True
                break

        if not section_found and current_section:
            if current_section == 'summary':
                resume_data[current_section] += ' ' + content
            elif current_section in ['skills', 'certifications']:
                items = re.split(r'[,•;\n]', content)
                resume_data[current_section].extend([item.strip() for item in items if item.strip()])
            else:
                lines = [e.strip() for e in content.split('\n') if e.strip()]
                resume_data[current_section].extend(lines)

    if resume_data['summary']:
        resume_data['summary'] = re.sub(
            r'^(profile\s*summary|summary|about|objective|profile|career\s*objectives)[:\s-]*',
            '',
            resume_data['summary'],
            flags=re.IGNORECASE
        ).strip()

    for section in list(resume_data.keys()):
        if isinstance(resume_data[section], (list, dict)) and not resume_data[section]:
            del resume_data[section]
        elif isinstance(resume_data[section], str) and not resume_data[section]:
            del resume_data[section]

    return resume_data


def create_output_doc(sections, original_filename):
    doc = Document()

    doc.add_heading('Parsed Resume Information', level=1)
    doc.add_paragraph(f"Original file: {original_filename}")
    doc.add_paragraph(f"Parsed on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    if 'photo' in sections and sections['photo']:
        try:
            for ext in ['.jpg', '.jpeg', '.png']:
                photo_path = os.path.join(app.config['UPLOAD_FOLDER'], f"photo_{sections['photo']}{ext}")
                if os.path.exists(photo_path):
                    doc.add_heading('Profile Photo', level=2)
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(photo_path, width=Inches(1.5))
                    break
        except Exception as e:
            app.logger.error(f"Error adding photo to document: {str(e)}")

    if 'contact' in sections:
        doc.add_heading('Contact Information', level=2)
        contact_table = doc.add_table(rows=0, cols=2)
        for field in ['name', 'email', 'phone', 'address', 'linkedin']:
            if sections['contact'].get(field):
                row = contact_table.add_row()
                row.cells[0].text = field.capitalize() + ':'
                row.cells[1].text = sections['contact'][field]

    if sections.get('summary'):
        doc.add_heading('Profile Summary', level=2)
        doc.add_paragraph(sections['summary'])

    for sec in ['skills', 'experience', 'education', 'certifications', 'projects', 'achievements']:
        if sections.get(sec):
            doc.add_heading(sec.capitalize(), level=2)
            for item in sections[sec]:
                doc.add_paragraph(item, style='List Bullet')

    output_stream = io.BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)
    return output_stream


def cleanup_uploads():
    """Clean up the uploads directory"""
    for filename in os.listdir(app.config['UPLOAD_FOLDER']):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print(f'Failed to delete {file_path}. Reason: {e}')


def process_zip_file(zip_path):
    """Process all DOCX files in a ZIP archive"""
    temp_folder = os.path.join(app.config['UPLOAD_FOLDER'], f'temp_{uuid.uuid4()}')
    os.makedirs(temp_folder, exist_ok=True)

    try:
        # Extract the ZIP file
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(temp_folder)

        # Process all DOCX files
        results = []
        for root, _, files in os.walk(temp_folder):
            for file in files:
                if file.lower().endswith('.docx'):
                    filepath = os.path.join(root, file)
                    try:
                        parsed_data = parse_resume(filepath)
                        output = create_output_doc(parsed_data, file)
                        results.append({
                            'filename': file,
                            'output': output
                        })
                    except Exception as e:
                        results.append({
                            'filename': file,
                            'error': str(e)
                        })

        # Create output ZIP if we found any DOCX files
        if not any('output' in result for result in results):
            raise ValueError("No valid .docx files found in the ZIP archive")

        output_zip_path = os.path.join(app.config['UPLOAD_FOLDER'], f'processed_resumes_{uuid.uuid4()}.zip')
        with zipfile.ZipFile(output_zip_path, 'w') as zipf:
            for result in results:
                if 'output' in result:
                    output_filename = f"parsed_{result['filename']}"
                    zipf.writestr(output_filename, result['output'].getvalue())

        return output_zip_path, results

    finally:
        # Clean up temporary files
        shutil.rmtree(temp_folder, ignore_errors=True)


@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        cleanup_uploads()  # Clean up previous files

        if 'zip_file' not in request.files or not request.files['zip_file']:
            flash('No file selected', 'error')
            return redirect(request.url)

        zip_file = request.files['zip_file']

        if zip_file.filename == '':
            flash('No file selected', 'error')
            return redirect(request.url)

        if not zip_file.filename.lower().endswith('.zip'):
            flash('Please upload a ZIP file', 'error')
            return redirect(request.url)

        try:
            # Save the uploaded ZIP file temporarily
            zip_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(zip_file.filename))
            zip_file.save(zip_path)

            # Process the ZIP file
            output_zip_path, results = process_zip_file(zip_path)

            # Count successful and failed files
            success_count = sum(1 for r in results if 'output' in r)
            error_count = sum(1 for r in results if 'error' in r)

            # Store results in session
            session['output_zip_path'] = output_zip_path
            session['processing_results'] = {
                'total': len(results),
                'success': success_count,
                'errors': error_count
            }

            # Clean up the uploaded ZIP
            os.remove(zip_path)

            flash(f'Processed {success_count} resumes successfully ({error_count} errors)', 'success')
            return redirect(url_for('download_results'))

        except Exception as e:
            flash(f'Error processing files: {str(e)}', 'error')
            return redirect(request.url)

    return render_template('upload.html')


@app.route('/download')
def download_results():
    if 'output_zip_path' not in session or not os.path.exists(session['output_zip_path']):
        flash('No processed files available. Please upload a ZIP file first.', 'error')
        return redirect(url_for('upload_file'))

    try:
        return send_file(
            session['output_zip_path'],
            as_attachment=True,
            download_name="processed_resumes.zip",
            mimetype='application/zip'
        )
    except Exception as e:
        flash(f'Error downloading files: {str(e)}', 'error')
        return redirect(url_for('upload_file'))
    finally:
        # Clean up the processed ZIP file after sending
        if 'output_zip_path' in session:
            try:
                os.remove(session['output_zip_path'])
            except:
                pass
            session.pop('output_zip_path', None)


@app.route('/photo/<photo_id>')
def serve_photo(photo_id):
    try:
        for ext in ['.jpg', '.jpeg', '.png']:
            photo_path = os.path.join(app.config['UPLOAD_FOLDER'], f"photo_{photo_id}{ext}")
            if os.path.exists(photo_path):
                return send_file(photo_path, mimetype=f'image/{ext[1:]}')
        return "Photo not found", 404
    except Exception as e:
        return f"Error loading photo: {str(e)}", 500


if __name__ == '__main__':
    app.run(debug=True)